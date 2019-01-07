# Python3
# coding: shift_jis
import configparser
import csv
import docx
import os
from datetime import datetime

# ��������������������������������������������������������������������
# ���Fshift_jis���Ŏg�p����ꍇ�́A
# 1.coding: shift_jis�Ƃ��邱�ƁI�I
# 2.�o�b�N�X���b�V��(�g�p�s��)��\�ɒu�����邱�ƁI�I
# 3.config.ini��encoding = shift_jis�Ƃ��邱�ƁI�I
# ��������������������������������������������������������������������
# �p�����[�^�ݒ�(config.ini���擾)
config = configparser.ConfigParser()

try:
    # config�̃t�@�C�����J��
    config.read('./config/config.ini')
    # �Z�N�V�������擾
    encoding_section = config['ENCODING']
    file_section = config['FILE_NAME']
    index_section = config['COLUMN_INDEX']

    # �p�����[�^�̓ǂݍ���
    # �G���R�[�h
    ENCODING = encoding_section.get('encoding')

    # �e�t�@�C����
    CSV_FILE_NAME = file_section.get('csv_file')
    DOCX_FILE_NAME = file_section.get('docx_file')

    # CSV������o������������̗�C���f�b�N�X(0�`)
    TITLE_COLUMN_INDEX = int(index_section.get('title_column'))
    QUESTION_COLUMN_INDEX = int(index_section.get('question_column'))
    FIRST_ANSWER_COLUMN_INDEX = int(index_section.get('first_answer_column'))
    LAST_ANSWER_COLUMN_INDEX = int(index_section.get('last_answer_column'))

# ini�t�@�C���I�[�v���A�Z�N�V�����擾�Ɏ��s�����ꍇ
# ���p�����[�^�ǂݍ��݃`�F�b�N�͎������Ă��܂���B
except:
    print("config.ini�����݂��Ȃ����Aconfig.ini���̃p�����[�^���s���ł��B")
    exit()

# --------------------------------------------------------
# Word�ɓ]�L����u�ŏI�񓚁v��Ώۍs���擾���郁�\�b�h
def get_last_answer(row):
    # �񓚗��1��̂݃p�����[�^�w�肵�Ă���ꍇ�A���̉񓚂��u�ŏI�񓚁v�Ƃ���
    if FIRST_ANSWER_COLUMN_INDEX == LAST_ANSWER_COLUMN_INDEX:
        return row[FIRST_ANSWER_COLUMN_INDEX]

    # �񓚂���납�珉��񓚂֌������A�ŏ��Ƀq�b�g�����񓚂��u�ŏI�񓚁v�Ƃ���
    for i in range(LAST_ANSWER_COLUMN_INDEX, FIRST_ANSWER_COLUMN_INDEX - 1, -1):
        if row[i] != "":
            return row[i]

    # ���ɑS�Ẳ񓚂��󔒂������ꍇ�A�󔒂�Ԃ�
    return ""

# --------------------------------------------------------
# csv�t�H���_�ɂ���CSV�t�@�C���I�[�v��
csv_path = "./csv/" + CSV_FILE_NAME
if not os.path.isfile(csv_path):
    print("csv�t�H���_��"+CSV_FILE_NAME+"�����݂��܂���B")
    exit()

csv_file = \
    open(csv_path, "r", encoding=ENCODING, errors="", newline="" )

# Word�t�@�C���V�K�쐬
doc = docx.Document()

# CSV�t�@�C���ǂݍ���
csv_rows = \
    csv.reader(csv_file, delimiter=",", doublequote=True, \
    lineterminator="\r\n", quotechar='"', skipinitialspace=True)

# 2�s�ڈȍ~��Word�t�@�C���ɓ]�L
for row in csv_rows:
    # 1�s�ڂ̃w�b�_�͓ǂݔ�΂�
    if csv_rows.line_num == 1:
        continue
    # ���o��
    doc.add_heading(row[TITLE_COLUMN_INDEX], 0)
    # ����
    doc.add_heading("����", 1)
    doc.add_paragraph(row[QUESTION_COLUMN_INDEX])
    doc.add_paragraph('\n')
    # ��(�ŏI�񓚂݂̂��Z�b�g����)
    doc.add_heading("��", 1)
    doc.add_paragraph(get_last_answer(row))
    # ���y�[�W
    doc.add_page_break()

# Word�t�@�C���̐V�K�쐬or�㏑��
docx_path = "./docx/"+DOCX_FILE_NAME
if os.path.isfile(docx_path):
    choice = input("�����t�@�C�������ɑ��݂��܂��B���l�[�����đޔ����܂����H\
        \n�͂� �� y(=n�ȊO) ������ �� n�F")
    # �����t�@�C�������l�[�����đޔ�
    if choice != "n":
        current_date = datetime.now().strftime('%y%m%d%H%M%S')
        os.rename(docx_path, str.rstrip(docx_path, ".docx") + current_date + ".docx")

# Word�t�@�C����docx�t�H���_�ɕۑ�
doc.save(docx_path)
print(DOCX_FILE_NAME+"���쐬���܂����B")
