# Python3
# utf-8
import configparser
import csv
import docx
# --------------------------------------------------------
# ▼仕様
# ・CSVファイルに含まれる「見出し」「質問」「回答」を
#   Wordファイルに転記します。
# ・CSVファイルは、当pyファイルと同一ディレクトリ内に存在すること、
#   パラメータ設定「config.ini」も同ディレクトリ内に存在すること、
#   Wordファイルも、同ディレクトリ内に作成or上書きされることを
#   想定しています。
# ・CSVファイル1行につき1つの質問、
#   質問に対する回答は1〜n回分あることを想定しています。
# ・Wordファイルに転記する回答は「最終回答」のみとします。
# ・1見出し/質問/回答を出力後に改ページします。
# ・「見出し」のスタイルはTitle固定、
#   「質問」「回答」(小見出し)のスタイルはHeading 1固定、
#   「質問」「回答」の内容自体はスタイル指定なしとします。
# ・Title、Heading 1の書式はWord側で適宜編集してください。
# ・config.iniにファイル名、列インデックスを指定してください。
# --------------------------------------------------------
# パラメータ設定(config.iniより取得)
config = configparser.ConfigParser()

try:
    # configのファイルを開く
    config.read('config.ini')
    # セクションを取得
    file_section = config['FILE_NAME']
    index_section = config['COLUMN_INDEX']

    # パラメータの読み込み
    # 各ファイル名
    CSV_FILE_NAME = file_section.get('csv_file')
    DOCX_FILE_NAME = file_section.get('docx_file')

    # CSVから取り出したい文字列の列インデックス(0〜)
    TITLE_COLUMN_INDEX = int(index_section.get('title_column'))
    QUESTION_COLUMN_INDEX = int(index_section.get('question_column'))
    FIRST_ANSWER_COLUMN_INDEX = int(index_section.get('first_answer_column'))
    LAST_ANSWER_COLUMN_INDEX = int(index_section.get('last_answer_column'))

# iniファイルオープン、セクション取得に失敗した場合
# ※パラメータ読み込みチェックは実装していません。
except:
    print("config.iniが存在しないか、config.ini内のパラメータが不正です。")
    exit()

# --------------------------------------------------------
# Wordに転記する「最終回答」を対象行より取得するメソッド
def get_last_answer(row):
    # 回答列を1列のみパラメータ指定している場合、その回答を「最終回答」とする
    if FIRST_ANSWER_COLUMN_INDEX == LAST_ANSWER_COLUMN_INDEX:
        return row[FIRST_ANSWER_COLUMN_INDEX]

    # 回答を後ろから初回回答へ検索し、最初にヒットした回答を「最終回答」とする
    for i in range(LAST_ANSWER_COLUMN_INDEX, FIRST_ANSWER_COLUMN_INDEX, -1):
        if row[i] != "":
            return row[i]

    # 仮に全ての回答が空白だった場合、空白を返す
    return ""

# --------------------------------------------------------
# カレントディレクトリにあるCSVファイルオープン
csv_file = \
    open("./" + CSV_FILE_NAME, "r", encoding="utf_8", errors="", newline="" )

# Wordファイル新規作成
doc = docx.Document()

# CSVファイル読み込み
csv_rows = \
    csv.reader(csv_file, delimiter=",", doublequote=True, \
    lineterminator="\r\n", quotechar='"', skipinitialspace=True)

# 1行目のヘッダは読み飛ばし
header = next(csv_rows)

# 2行目以降をWordファイルに転記
for row in csv_rows:
    # 見出し
    doc.add_heading(row[TITLE_COLUMN_INDEX], 0)
    # 質問
    doc.add_heading("質問", 1)
    doc.add_paragraph(row[QUESTION_COLUMN_INDEX])
    doc.add_paragraph('\n')
    # 回答(最終回答のみをセットする)
    doc.add_heading("回答", 1)
    doc.add_paragraph(get_last_answer(row))
    # 改ページ
    doc.add_page_break()

# Wordファイルをカレントディレクトリに保存
doc.save(DOCX_FILE_NAME)
